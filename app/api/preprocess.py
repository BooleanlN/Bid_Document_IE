import docx
import PyPDF2
import pdfplumber

import os
import re

class ClipFile:
    def __init__(self,file_source):
        self._file_source = file_source
    def handle_pdf(self,filename,destination="../source/extract/"):
        """
        处理原始招标文件(pdf格式)
        :param filename: 待处理文件名
        :param destination: 生成文件目标文件夹
        :return:
        """
        pdf_obj = PyPDF2.PdfFileReader(self._file_source + filename)
        doc_info = pdf_obj.documentInfo
        print(doc_info)

        #     new_meta = {
        #         "Author":doc_info.dc_creator,
        #         "CreateDate":doc_info.xmp_createDate,
        #         "ModifyDate":doc_info.xmp_modifyDate

        #     }
        outlines = pdf_obj.getOutlines()
        pattern = '.*(技术|参数|要求).*'
        pdf_out = PyPDF2.PdfFileWriter()
        pdf_out.addMetadata(doc_info)
        sections = []
        for outline in outlines:
            if isinstance(outline, list):
                continue
            #         print(outline)
            #             print(page.extractText())
            if re.match("(第.*章|第.*部分)", outline.title):
                # print(outline.title)
                sections.append(outline)
        for index, section in enumerate(sections):
            if re.match(".*(招标公告|招标邀请书).*", section.title):
                pageNumberStart = pdf_obj.getDestinationPageNumber(section)
                if isinstance(outlines[index + 1], list):
                    pageNumberEnd = pdf_obj.getDestinationPageNumber(sections[index + 1][0])
                else:
                    pageNumberEnd = pdf_obj.getDestinationPageNumber(sections[index + 1])
                # print(pageNumberStart,pageNumberEnd)
                bid_date = self.get_metadata(filename, pageNumberStart, pageNumberEnd)
                cc = {}
                cc["BidDate"] = bid_date
                pdf_out.addMetadata(cc)
            if re.match(pattern, section.title):
                # print(section.title)
                pageNumberStart = pdf_obj.getDestinationPageNumber(section)
                if isinstance(outlines[index + 1], list):
                    pageNumberEnd = pdf_obj.getDestinationPageNumber(sections[index + 1][0])
                else:
                    pageNumberEnd = pdf_obj.getDestinationPageNumber(sections[index + 1])
                # print(pageNumberStart,pageNumberEnd)
                for page in range(pageNumberStart, pageNumberEnd):
                    pdf_out.addPage(pdf_obj.getPage(page))
        with open(destination + filename, "wb") as out_file:
            try:
                pdf_out.write(out_file)
            except UnicodeEncodeError:
                print("编码错误")
                pass
    def handle_docx(self,filename,destination="../source/extract/"):
        """
        截取docx格式文件
        :param filename:文件名
        :param destination:目标文件夹
        :return:
        """
        doc_obj = docx.Document(self._file_source + filename)
        paragraphs = doc_obj.paragraphs
        #     for table in doc_obj.tables:
        #         print(table)
        sections = []
        pattern = '.*(技术|参数|要求).*'
        doc_out = docx.Document()
        for index, paragraph in enumerate(paragraphs):
            # print(paragraph.style.name)
            if "Heading" in paragraph.style.name:
                sections.append((index, paragraph))
        for order, (index, section) in enumerate(sections):
            if re.match(pattern, section.text):
                #             print(paragraphs[index].text)
                # doc_out.add_paragraph(paragraphs)
                for para in paragraphs[index:sections[order + 1][0]]:
                    doc_out.add_paragraph(para.text, para.style)
        doc_out.save(destination+ filename)

    def get_metadata(self, filename, pageNumberStart, pageNumberEnd):
        """
        获取招标时间
        :param filename:
        :param pageNumberStart:
        :param pageNumberEnd:
        :return:
        """
        origin = pdfplumber.open(self._file_source + filename)
        pages = origin.pages[pageNumberStart:pageNumberEnd]
        pattern = ".*(开标时间|投标截止时间).*"
        res = []
        for page in pages:
            texts = page.extract_text().split('\n')
            for index, text in enumerate(texts):
                #             print(text)
                if re.match(pattern, text):
                    res.append(text)
                    res.append(texts[index + 1])
        return "".join(res)
    def run(self):
        """
        执行
        :return:
        """
        for _, _, filenames in os.walk(self._file_source):
            for filename in filenames:
                print("{} is dealing with".format(filename))
                if filename.split(".")[-1] == 'pdf':
                    # continue
                    self.handle_pdf(filename)
                elif filename.split(".")[-1] == 'docx':
                    self.handle_docx(filename)

